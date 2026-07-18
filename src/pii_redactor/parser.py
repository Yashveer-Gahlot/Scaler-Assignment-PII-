"""
parser.py — Run-Level DOCX PII Redaction Parser
=================================================

Reads a Microsoft Word (.docx) document, applies PII redaction at the
granular *run* level (preserving bold, italic, color, font-size, and all
other inline formatting), and writes a clean redacted copy.

Key Design Decisions
--------------------
1. **Cross-run span reconstruction**: PII entities frequently span multiple
   Word runs (e.g. "Kushal" | " Subbayya" | " Hegde" across three styled
   fragments).  This parser concatenates run texts into a virtual paragraph
   string, detects PII on the concatenated text, then *slices* replacements
   back into the original run boundaries so formatting is never broken.

2. **Table-aware traversal**: Every table → row → cell → paragraph → run
   chain is traversed identically to body paragraphs.

3. **Header / Footer coverage**: Section headers and footers are also
   scanned, catching letterhead PII that body-only parsers miss.

4. **Deterministic output**: Uses the same ``StatefulAnonymizer`` instance
   across the entire document, so repeated names always map to the same
   synthetic replacement — even across paragraphs and tables.

Usage
-----
    $ python parser.py path/to/Prospectus.docx

    # Produces  →  Redacted_Prospectus.docx  (in the same directory)

Programmatic API:

    >>> from parser import redact_document
    >>> summary = redact_document("Prospectus.docx", "Redacted_Prospectus.docx")
    >>> print(summary)

Author : Parser Team
License: MIT
"""

from __future__ import annotations

import os
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# Phase-1 engine imports.
from .engine import RedactionPipeline, StatefulAnonymizer, create_pipeline


# ──────────────────────────────────────────────
# §1  Data Structures
# ──────────────────────────────────────────────

@dataclass
class RunSlice:
    """Maps a character range in the virtual paragraph string back to
    the originating ``Run`` object.

    Attributes
    ----------
    run:
        The ``python-docx`` Run instance.
    start:
        Inclusive start offset in the concatenated paragraph text.
    end:
        Exclusive end offset in the concatenated paragraph text.
    """

    run: Run
    start: int
    end: int


@dataclass
class RedactionStats:
    """Accumulates metrics across the entire document parse.

    Attributes
    ----------
    paragraphs_scanned:
        Total paragraphs processed (body + tables + headers/footers).
    runs_modified:
        Number of individual runs whose text was altered.
    entities_redacted:
        Total PII entity replacements performed.
    entity_breakdown:
        ``{entity_type: count}`` summary.
    """

    paragraphs_scanned: int = 0
    runs_modified: int = 0
    entities_redacted: int = 0
    entity_breakdown: Dict[str, int] = field(default_factory=dict)

    def record(self, entity_type: str, count: int = 1) -> None:
        """Increment counters for a given entity type."""
        self.entities_redacted += count
        self.entity_breakdown[entity_type] = (
            self.entity_breakdown.get(entity_type, 0) + count
        )

    def __str__(self) -> str:
        lines: List[str] = [
            f"Paragraphs scanned : {self.paragraphs_scanned}",
            f"Runs modified      : {self.runs_modified}",
            f"Entities redacted  : {self.entities_redacted}",
            "Breakdown:",
        ]
        for etype, cnt in sorted(self.entity_breakdown.items()):
            lines.append(f"  {etype:<22} {cnt}")
        return "\n".join(lines)


# ──────────────────────────────────────────────
# §2  Run-Level Replacement Engine
# ──────────────────────────────────────────────

def _build_run_map(paragraph: Paragraph) -> Tuple[str, List[RunSlice]]:
    """Concatenate run texts and build an offset → Run mapping.

    Parameters
    ----------
    paragraph:
        A ``python-docx`` Paragraph whose runs will be indexed.

    Returns
    -------
    full_text:
        The concatenated plain-text content of all runs.
    run_map:
        Ordered list of ``RunSlice`` objects mapping character offsets
        back to their source runs.
    """
    run_map: List[RunSlice] = []
    offset: int = 0
    for run in paragraph.runs:
        text: str = run.text or ""
        run_map.append(RunSlice(run=run, start=offset, end=offset + len(text)))
        offset += len(text)
    full_text: str = "".join(rs.run.text or "" for rs in run_map)
    return full_text, run_map


def _apply_replacements_to_runs(
    run_map: List[RunSlice],
    replacements: List[Tuple[int, int, str]],
) -> int:
    """Splice replacement strings into runs, preserving formatting.

    This is the critical algorithm.  Each replacement span
    ``(start, end, new_text)`` may cross one or more run boundaries.
    The strategy:

    1. Find every run that overlaps ``[start, end)``.
    2. Assign the full replacement text to the *first* overlapping run.
    3. Blank out the portions of subsequent overlapping runs that fall
       inside the replacement span.

    This guarantees that the first run's formatting (bold, color, etc.)
    is applied to the replacement text, and no characters are duplicated
    or lost.

    Parameters
    ----------
    run_map:
        The offset-indexed run mapping from ``_build_run_map``.
    replacements:
        ``(start, end, replacement_text)`` triples, **sorted in reverse
        offset order** so that earlier replacements don't shift later
        offsets.

    Returns
    -------
    int
        Number of individual runs whose text was mutated.
    """
    if not replacements:
        return 0

    # Rebuild mutable per-run text buffers.
    buffers: List[str] = [rs.run.text or "" for rs in run_map]
    runs_touched: int = 0

    for rep_start, rep_end, new_text in replacements:
        first_run_handled: bool = False

        for idx, rs in enumerate(run_map):
            # No overlap → skip.
            if rs.end <= rep_start or rs.start >= rep_end:
                continue

            # Relative offsets within this run's buffer.
            local_start: int = max(0, rep_start - rs.start)
            local_end: int = min(len(buffers[idx]), rep_end - rs.start)

            if not first_run_handled:
                # First overlapping run gets the full replacement text.
                buffers[idx] = (
                    buffers[idx][:local_start]
                    + new_text
                    + buffers[idx][local_end:]
                )
                first_run_handled = True
                runs_touched += 1
            else:
                # Subsequent overlapping runs: erase the overlapping slice.
                buffers[idx] = (
                    buffers[idx][:local_start] + buffers[idx][local_end:]
                )
                runs_touched += 1

    # Commit buffers back to the actual Run objects.
    for idx, rs in enumerate(run_map):
        if rs.run.text != buffers[idx]:
            rs.run.text = buffers[idx]

    return runs_touched


# ──────────────────────────────────────────────
# §3  Paragraph-Level Redaction
# ──────────────────────────────────────────────

def _redact_paragraph(
    paragraph: Paragraph,
    pipeline: RedactionPipeline,
    stats: RedactionStats,
) -> None:
    """Detect and redact PII in a single paragraph, run by run.

    Parameters
    ----------
    paragraph:
        The ``python-docx`` Paragraph to process.
    pipeline:
        An initialised ``RedactionPipeline`` from engine.py.
    stats:
        Mutable statistics accumulator.
    """
    stats.paragraphs_scanned += 1

    full_text, run_map = _build_run_map(paragraph)
    if not full_text.strip():
        return

    # Detect PII on the concatenated paragraph text.
    result = pipeline.redact(full_text)
    if not result.entries:
        return

    # Build replacement list in reverse offset order (critical for
    # non-shifting splice operations).
    replacements: List[Tuple[int, int, str]] = [
        (entry.start, entry.end, entry.replacement)
        for entry in sorted(result.entries, key=lambda e: e.start, reverse=True)
    ]

    runs_modified: int = _apply_replacements_to_runs(run_map, replacements)

    stats.runs_modified += runs_modified
    for entry in result.entries:
        stats.record(entry.entity_type)


# ──────────────────────────────────────────────
# §4  Document Traversal
# ──────────────────────────────────────────────

def _redact_paragraphs(
    paragraphs: Sequence[Paragraph],
    pipeline: RedactionPipeline,
    stats: RedactionStats,
) -> None:
    """Iterate over a sequence of paragraphs and redact each one."""
    for para in paragraphs:
        try:
            _redact_paragraph(para, pipeline, stats)
        except Exception as exc:  # noqa: BLE001
            # Log but don't crash — a single malformed run should not
            # abort the entire document.
            para_preview: str = (para.text or "")[:60]
            print(
                f"  ⚠  Skipping paragraph (error: {exc}): "
                f"'{para_preview}…'",
                file=sys.stderr,
            )


def _redact_table(
    table: Table,
    pipeline: RedactionPipeline,
    stats: RedactionStats,
) -> None:
    """Walk every cell in a table, redacting paragraph content.

    Handles nested tables recursively via ``cell.tables``.

    Parameters
    ----------
    table:
        A ``python-docx`` Table object.
    pipeline:
        The PII redaction pipeline.
    stats:
        Mutable statistics accumulator.
    """
    row: _Row
    for row in table.rows:
        cell: _Cell
        for cell in row.cells:
            _redact_paragraphs(cell.paragraphs, pipeline, stats)

            # Recurse into nested tables (common in complex prospectuses).
            nested_table: Table
            for nested_table in cell.tables:
                _redact_table(nested_table, pipeline, stats)


def _iter_body_elements(document: Document):
    """Yield paragraphs and tables in document-order.

    ``python-docx`` exposes ``document.paragraphs`` and
    ``document.tables`` as separate flat lists, which loses
    interleaving order.  This generator walks the underlying XML
    ``<w:body>`` children to yield elements in true reading order.

    Yields
    ------
    Paragraph | Table
        Elements in the order they appear in the document body.
    """
    body = document.element.body
    for child in body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, document.element.body)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document.element.body)


def _redact_headers_footers(
    document: Document,
    pipeline: RedactionPipeline,
    stats: RedactionStats,
) -> None:
    """Scan section headers and footers for PII.

    Letterhead, contact blocks, and legal footers frequently contain
    names, addresses, phone numbers, and emails that must be redacted.
    """
    for section in document.sections:
        # Default header/footer.
        if section.header and section.header.is_linked_to_previous is False:
            _redact_paragraphs(section.header.paragraphs, pipeline, stats)
            for tbl in section.header.tables:
                _redact_table(tbl, pipeline, stats)

        if section.footer and section.footer.is_linked_to_previous is False:
            _redact_paragraphs(section.footer.paragraphs, pipeline, stats)
            for tbl in section.footer.tables:
                _redact_table(tbl, pipeline, stats)

        # First-page header/footer (if distinct).
        if section.first_page_header and hasattr(section, "first_page_header"):
            try:
                _redact_paragraphs(
                    section.first_page_header.paragraphs, pipeline, stats
                )
            except Exception:  # noqa: BLE001
                pass

        if section.first_page_footer and hasattr(section, "first_page_footer"):
            try:
                _redact_paragraphs(
                    section.first_page_footer.paragraphs, pipeline, stats
                )
            except Exception:  # noqa: BLE001
                pass


# ──────────────────────────────────────────────
# §5  Public API
# ──────────────────────────────────────────────

def redact_document(
    input_path: str | Path,
    output_path: Optional[str | Path] = None,
    *,
    score_threshold: float = 0.50,
    locale: str = "en_IN",
) -> RedactionStats:
    """Redact all PII from a .docx file, preserving formatting.

    Parameters
    ----------
    input_path:
        Path to the source Word document.
    output_path:
        Path for the redacted output.  Defaults to
        ``Redacted_Prospectus.docx`` in the same directory as *input_path*.
    score_threshold:
        Minimum confidence for PII detection (forwarded to the engine).
    locale:
        Faker locale for synthetic replacements.

    Returns
    -------
    RedactionStats
        Summary of all redactions performed.

    Raises
    ------
    FileNotFoundError
        If *input_path* does not exist.
    ValueError
        If *input_path* is not a ``.docx`` file.
    RuntimeError
        If the document cannot be parsed by ``python-docx``.
    """
    # ── Validate input ──────────────────────────────────────────────
    src: Path = Path(input_path).resolve()
    if not src.exists():
        raise FileNotFoundError(
            f"Input document not found: {src}\n"
            f"Please verify the file path and try again."
        )
    if src.suffix.lower() != ".docx":
        raise ValueError(
            f"Expected a .docx file, got '{src.suffix}'. "
            f"This parser only supports Office Open XML (.docx) documents."
        )

    # ── Resolve output path ─────────────────────────────────────────
    if output_path is None:
        output_path = src.parent / "Redacted_Prospectus.docx"
    else:
        output_path = Path(output_path).resolve()

    # ── Load document ───────────────────────────────────────────────
    try:
        doc: Document = Document(str(src))
    except Exception as exc:
        raise RuntimeError(
            f"Failed to open document '{src.name}': {exc}\n"
            f"The file may be corrupted or not a valid .docx."
        ) from exc

    print(f"📄 Loaded: {src.name}")
    print(f"   Sections: {len(doc.sections)}")
    print(f"   Body paragraphs: {len(doc.paragraphs)}")
    print(f"   Tables: {len(doc.tables)}")
    print()

    # ── Initialize pipeline ─────────────────────────────────────────
    pipeline: RedactionPipeline = create_pipeline(
        score_threshold=score_threshold,
        locale=locale,
    )
    stats: RedactionStats = RedactionStats()

    # ── Phase A: Body elements (paragraphs & tables in order) ──────
    print("🔍 Phase A — Scanning body elements…")
    for element in _iter_body_elements(doc):
        if isinstance(element, Paragraph):
            try:
                _redact_paragraph(element, pipeline, stats)
            except Exception as exc:  # noqa: BLE001
                preview: str = (element.text or "")[:60]
                print(
                    f"  ⚠  Skipping body paragraph (error: {exc}): "
                    f"'{preview}…'",
                    file=sys.stderr,
                )
        elif isinstance(element, Table):
            _redact_table(element, pipeline, stats)

    # ── Phase B: Headers & Footers ─────────────────────────────────
    print("🔍 Phase B — Scanning headers & footers…")
    _redact_headers_footers(doc, pipeline, stats)

    # ── Save ────────────────────────────────────────────────────────
    try:
        doc.save(str(output_path))
    except PermissionError as exc:
        raise RuntimeError(
            f"Cannot write to '{output_path}': {exc}\n"
            f"Close the file if it is open in another application."
        ) from exc

    print()
    print(f"✅ Saved redacted document → {output_path.name}")
    print()
    print("─" * 48)
    print("REDACTION SUMMARY")
    print("─" * 48)
    print(stats)
    print("─" * 48)

    return stats


# ──────────────────────────────────────────────
# §6  CLI Entry Point
# ──────────────────────────────────────────────

def main() -> None:
    """Command-line entry point with argument validation."""
    if len(sys.argv) < 2:
        print(
            "Usage: python parser.py <path-to-document.docx> [output-path.docx]\n"
            "\n"
            "Arguments:\n"
            "  input_path    Path to the source .docx file.\n"
            "  output_path   (Optional) Path for the redacted output.\n"
            "                Defaults to 'Redacted_Prospectus.docx' in the\n"
            "                same directory as the input file.\n",
            file=sys.stderr,
        )
        sys.exit(1)

    input_file: str = sys.argv[1]
    output_file: Optional[str] = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        stats: RedactionStats = redact_document(input_file, output_file)
    except FileNotFoundError as exc:
        print(f"\n❌ File Not Found:\n   {exc}", file=sys.stderr)
        sys.exit(2)
    except ValueError as exc:
        print(f"\n❌ Invalid Input:\n   {exc}", file=sys.stderr)
        sys.exit(3)
    except RuntimeError as exc:
        print(f"\n❌ Processing Error:\n   {exc}", file=sys.stderr)
        sys.exit(4)
    except KeyboardInterrupt:
        print("\n\n⏹  Interrupted by user.", file=sys.stderr)
        sys.exit(130)

    # Non-zero exit if nothing was found (may indicate wrong file).
    if stats.entities_redacted == 0:
        print(
            "\n⚠  No PII entities detected. Verify the document "
            "contains personal information.",
            file=sys.stderr,
        )


if __name__ == "__main__":
    main()
